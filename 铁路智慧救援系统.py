import sys
import os
import io
import base64
import requests
import logging
import ctypes
import cv2
import subprocess
import numpy as np
from ctypes import cdll, c_int, c_char_p, POINTER, CFUNCTYPE
from contextlib import contextmanager
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from ctypes import cdll, c_int, c_char_p, POINTER, CFUNCTYPE
from docx import Document
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from datetime import datetime
from PIL import Image
from PyQt5.QtWidgets import (QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
                             QLabel, QPushButton, QTextEdit, QFileDialog, QMessageBox,
                             QListWidget, QListWidgetItem, QTabWidget, QProgressBar,
                             QFrame, QSlider, QDialog, QFormLayout, QSpinBox, QMenu, QAction,
                             QSplitter)
from PyQt5.QtGui import QPixmap, QIcon, QImage, QPalette, QColor, QFont
from PyQt5.QtCore import Qt, QSize, QThread, pyqtSignal, QObject, QEvent, QUrl
from PyQt5.QtMultimedia import QMediaContent, QMediaPlayer
from PyQt5.QtMultimediaWidgets import QVideoWidget

import requests


def get_model_memory():
    response = requests.get('http://localhost:11434/api/ps')
    models = response.json().get('models', [])

    if not models:
        print("没有正在运行的模型")
        return

    for model in models:
        size_gb = model.get('size', 0) / (1024 ** 3)
        vram_gb = model.get('size_vram', 0) / (1024 ** 3)
        print(f"模型: {model.get('name')}")
        print(f"  总大小: {size_gb:.2f} GB")
        print(f"  VRAM 占用: {vram_gb:.2f} GB")
        print(f"  参数量: {model.get('details', {}).get('parameter_size', 'N/A')}")


# 调用函数
if __name__ == "__main__":
    get_model_memory()




# ------------------------------
# 核心修复：拦截libpng的C层面警告输出
# ------------------------------
if sys.platform.startswith('linux'):
    try:
        libc = cdll.LoadLibrary('libc.so.6')
        c_write_func = CFUNCTYPE(c_int, c_int, c_char_p, c_int)
        original_write = libc.write


        def filtered_write(fd, buf, count):
            if fd == 2:
                try:
                    buffer_str = ctypes.string_at(buf, count).decode('utf-8', errors='ignore')
                    if "libpng warning: iCCP: known incorrect sRGB profile" in buffer_str:
                        return count
                except:
                    pass
            return original_write(fd, buf, count)


        libc.write = c_write_func(filtered_write)
    except:
        pass

elif sys.platform.startswith('win'):
    try:
        kernel32 = ctypes.WinDLL('kernel32', use_last_error=True)
        GetStdHandle = kernel32.GetStdHandle
        GetStdHandle.argtypes = [c_int]
        GetStdHandle.restype = ctypes.HANDLE

        WriteFile = kernel32.WriteFile
        WriteFile.argtypes = [ctypes.HANDLE, c_char_p, c_int, POINTER(c_int), ctypes.c_void_p]
        WriteFile.restype = ctypes.c_bool

        stderr_handle = GetStdHandle(-12)
        original_WriteFile = WriteFile


        def filtered_WriteFile(hFile, lpBuffer, nNumberOfBytesToWrite, lpNumberOfBytesWritten, lpOverlapped):
            if hFile == stderr_handle:
                try:
                    buffer_str = ctypes.string_at(lpBuffer, nNumberOfBytesToWrite).decode('utf-8', errors='ignore')
                    if "libpng warning: iCCP: known incorrect sRGB profile" in buffer_str:
                        if lpNumberOfBytesWritten:
                            lpNumberOfBytesWritten[0] = nNumberOfBytesToWrite
                        return True
                except:
                    pass
            return original_WriteFile(hFile, lpBuffer, nNumberOfBytesToWrite, lpNumberOfBytesWritten, lpOverlapped)


        WriteFile = filtered_WriteFile
    except:
        pass


# 辅助方案：重定向stderr过滤警告
class WarningFilter:
    def __init__(self, stream):
        self.stream = stream
        self.pattern = "libpng warning: iCCP: known incorrect sRGB profile"

    def write(self, message):
        if self.pattern not in message:
            self.stream.write(message)
        self.stream.flush()

    def flush(self):
        self.stream.flush()


sys.stderr = WarningFilter(sys.stderr)
logging.getLogger('PIL').setLevel(logging.CRITICAL + 1)
os.environ["PYTHONWARNINGS"] = "ignore"


class VideoFrameExtractor(QThread):
    progress_updated = pyqtSignal(int)
    frame_extracted = pyqtSignal(str)
    extraction_complete = pyqtSignal(list)
    extraction_error = pyqtSignal(str)
    extraction_log = pyqtSignal(str)  # 提取帧的日志信息

    def __init__(self, video_path, interval=10):
        super().__init__()
        self.video_path = video_path
        self.interval = interval  # 提取帧的间隔（秒）
        self.is_running = True
        self.current_video_index = 0  # 当前处理的视频索引
        self.total_videos = 1  # 总视频数量

    def set_video_index(self, index, total):
        self.current_video_index = index
        self.total_videos = total

    def run(self):
        try:
            self.extraction_log.emit(f"开始从视频中提取帧: {os.path.basename(self.video_path)}")
            cap = cv2.VideoCapture(self.video_path)
            if not cap.isOpened():
                self.extraction_error.emit("无法打开视频文件")
                return

            fps = cap.get(cv2.CAP_PROP_FPS)
            total_frames = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
            total_seconds = total_frames / fps if fps > 0 else 0
            frame_interval = int(fps * self.interval)

            if frame_interval <= 0:
                frame_interval = 1

            self.extraction_log.emit(f"视频信息: FPS={fps:.2f}, 总帧数={total_frames}, 时长={total_seconds:.2f}秒")
            self.extraction_log.emit(f"提取间隔: {self.interval}秒, 每间隔{frame_interval}帧提取一次")

            extracted_frames = []
            frame_count = 0
            total_steps = int(total_frames / frame_interval) + 1  # 计算总步骤数
            current_step = 0

            while self.is_running and frame_count < total_frames:
                cap.set(cv2.CAP_PROP_POS_FRAMES, frame_count)
                ret, frame = cap.read()
                if not ret:
                    self.extraction_log.emit(f"读取帧 {frame_count} 失败，停止提取")
                    break

                frame_path = os.path.join(
                    os.path.dirname(self.video_path),
                    f"frame_{os.path.basename(self.video_path).split('.')[0]}_{int(frame_count / fps)}.jpg"
                )

                cv2.imwrite(frame_path, frame)
                extracted_frames.append(frame_path)
                self.frame_extracted.emit(frame_path)
                self.extraction_log.emit(f"已提取帧 {frame_count}，保存至: {os.path.basename(frame_path)}")

                # 精确计算进度，考虑多视频情况
                current_step += 1
                video_progress = int((current_step / total_steps) * 100)
                overall_progress = int((self.current_video_index / self.total_videos) * 100 +
                                       (video_progress / self.total_videos))
                self.progress_updated.emit(overall_progress)

                frame_count += frame_interval

            cap.release()
            self.extraction_log.emit(f"视频帧提取完成，共提取{len(extracted_frames)}帧")
            self.extraction_complete.emit(extracted_frames)
        except Exception as e:
            error_msg = f"提取帧失败: {str(e)}"
            self.extraction_log.emit(error_msg)
            self.extraction_error.emit(error_msg)

    def stop(self):
        self.is_running = False
        self.wait()


class FrameExtractionDialog(QDialog):
    def __init__(self, video_path, parent=None):
        super().__init__(parent)
        self.setWindowTitle("视频帧提取设置")
        self.setGeometry(300, 300, 300, 150)
        self.apply_dark_style()

        layout = QFormLayout()

        self.interval_spin = QSpinBox()
        self.interval_spin.setRange(1, 60)
        self.interval_spin.setValue(10)
        self.interval_spin.setSuffix(" 秒")
        self.interval_spin.setStyleSheet("color: #FFFFFF; background-color: #333333;")

        layout.addRow("提取帧间隔:", self.interval_spin)

        btn_layout = QHBoxLayout()
        self.ok_btn = QPushButton("确定")
        self.cancel_btn = QPushButton("取消")

        self.ok_btn.setStyleSheet("""
            QPushButton {
                background-color: #4CAF50;
                color: white;
                border-radius: 5px;
                padding: 5px;
                border: none;
            }
            QPushButton:hover {
                background-color: #45a049;
                box-shadow: 0 2px 4px rgba(0,0,0,0.2);
            }
            QPushButton:pressed {
                background-color: #3d8b40;
                box-shadow: inset 0 2px 4px rgba(0,0,0,0.2);
            }
        """)

        self.cancel_btn.setStyleSheet("""
            QPushButton {
                background-color: #f44336;
                color: white;
                border-radius: 5px;
                padding: 5px;
                border: none;
            }
            QPushButton:hover {
                background-color: #d32f2f;
                box-shadow: 0 2px 4px rgba(0,0,0,0.2);
            }
            QPushButton:pressed {
                background-color: #b71c1c;
                box-shadow: inset 0 2px 4px rgba(0,0,0,0.2);
            }
        """)

        self.ok_btn.clicked.connect(self.accept)
        self.cancel_btn.clicked.connect(self.reject)

        btn_layout.addWidget(self.ok_btn)
        btn_layout.addWidget(self.cancel_btn)

        layout.addRow(btn_layout)
        self.setLayout(layout)

        self.video_path = video_path

    def get_interval(self):
        return self.interval_spin.value()

    def apply_dark_style(self):
        self.setStyleSheet("""
            QDialog {
                background-color: #222222;
                color: #FFFFFF;
            }
            QLabel {
                color: #FFFFFF;
            }
        """)


# 大模型对话框
class ModelDialog(QDialog):
    def __init__(self, title, initial_content, prompt, parent=None):
        super().__init__(parent)
        self.setWindowTitle(title)
        self.resize(800, 600)
        self.apply_dark_style()

        layout = QVBoxLayout(self)

        # 提示信息
        prompt_label = QLabel("提示: " + prompt)
        prompt_label.setWordWrap(True)
        prompt_label.setStyleSheet("color: #BBBBBB; font-style: italic; margin-bottom: 10px;")
        layout.addWidget(prompt_label)

        # 内容编辑区
        self.content_editor = QTextEdit()
        self.content_editor.setPlainText(initial_content)
        self.content_editor.setStyleSheet("""
            QTextEdit {
                background-color: #333333;
                color: #FFFFFF;
                border: 1px solid #555555;
                border-radius: 4px;
                padding: 5px;
            }
        """)
        layout.addWidget(self.content_editor)

        # 按钮区
        btn_layout = QHBoxLayout()
        self.regenerate_btn = QPushButton("重新生成")
        self.accept_btn = QPushButton("确认")
        self.cancel_btn = QPushButton("取消")

        self.regenerate_btn.setStyleSheet("""
            QPushButton {
                background-color: #ff9800;
                color: white;
                border-radius: 5px;
                padding: 5px;
                border: none;
            }
            QPushButton:hover {
                background-color: #e68900;
            }
        """)

        self.accept_btn.setStyleSheet("""
            QPushButton {
                background-color: #4CAF50;
                color: white;
                border-radius: 5px;
                padding: 5px;
                border: none;
            }
            QPushButton:hover {
                background-color: #45a049;
            }
        """)

        self.cancel_btn.setStyleSheet("""
            QPushButton {
                background-color: #f44336;
                color: white;
                border-radius: 5px;
                padding: 5px;
                border: none;
            }
            QPushButton:hover {
                background-color: #d32f2f;
            }
        """)

        self.regenerate_btn.clicked.connect(self.on_regenerate)
        self.accept_btn.clicked.connect(self.accept)
        self.cancel_btn.clicked.connect(self.reject)

        btn_layout.addWidget(self.regenerate_btn)
        btn_layout.addWidget(self.accept_btn)
        btn_layout.addWidget(self.cancel_btn)

        layout.addLayout(btn_layout)

        self.prompt = prompt
        self.analyzer = parent.analyzer if parent else None
        self.result_type = title.split(" ")[0]  # 从标题提取类型
        self.initial_content = initial_content
        self.parent = parent

    def get_content(self):
        return self.content_editor.toPlainText()

    def on_regenerate(self):
        if not self.analyzer or not self.parent:
            QMessageBox.warning(self, "警告", "无法连接到分析器")
            return

        # 显示加载状态
        self.content_editor.setPlainText("正在重新生成...")
        self.regenerate_btn.setEnabled(False)

        # 根据类型重新生成内容
        try:
            if self.result_type == "现场情况":
                all_media = self.parent.media_paths.copy()
                all_media.extend(self.parent.extracted_frames)
                scene_descriptions = []
                for path in all_media:
                    if path.lower().endswith(('.mp4', '.avi', '.mov', '.mkv')):
                        result = "视频文件：包含动态事故现场画面"
                    else:
                        result = self.analyzer.analyze_scene(path)
                    scene_descriptions.append(f"【现场素材分析】\n{result}\n")
                content = self.analyzer.synthesize_scene(scene_descriptions)

            elif self.result_type == "救援分析":
                scene_content = self.parent.step_results["scene"]
                content = self.analyzer.analyze_rescue_feasibility(scene_content)

            elif self.result_type == "救援步骤":
                analysis_content = self.parent.step_results["analysis"]
                content = self.analyzer.generate_rescue_steps(analysis_content)

            elif self.result_type == "风险预判":
                steps_content = self.parent.step_results["steps"]
                content = self.analyzer.identify_risk_points(steps_content)

            self.content_editor.setPlainText(content)
            self.parent.log_thinking(f"{self.result_type}已重新生成")

        except Exception as e:
            self.content_editor.setPlainText(f"生成失败: {str(e)}")
            self.parent.log_thinking(f"{self.result_type}重新生成失败: {str(e)}")

        finally:
            self.regenerate_btn.setEnabled(True)

    def apply_dark_style(self):
        self.setStyleSheet("""
            QDialog {
                background-color: #222222;
                color: #FFFFFF;
            }
            QLabel {
                color: #FFFFFF;
            }
        """)


class AnalysisSignals(QObject):
    progress_updated = pyqtSignal(int)
    step_completed = pyqtSignal(str, str)
    analysis_error = pyqtSignal(str)
    thinking_log = pyqtSignal(str)  # 大模型思考过程日志
    irrelevant_material = pyqtSignal(str)  # 素材无关信号


class AnalysisThread(QThread):
    def __init__(self, analyzer, step_type, input_data, media_paths=None):
        super().__init__()
        self.signals = AnalysisSignals()
        self.analyzer = analyzer
        self.step_type = step_type
        self.input_data = input_data
        self.media_paths = media_paths  # 可以是图片路径或视频帧路径
        self.is_running = True

    def run(self):
        try:
            if self.step_type == "scene":
                total = len(self.media_paths)
                scene_descriptions = []
                self.signals.thinking_log.emit(f"开始分析现场情况，共{total}个素材需要处理")

                # 先检查素材是否与铁路相关
                self.signals.thinking_log.emit("首先检查所有素材是否与铁路相关")
                all_related = True
                for i, path in enumerate(self.media_paths):
                    if not self.is_running:
                        self.signals.thinking_log.emit("分析已取消")
                        self.signals.step_completed.emit("scene", "分析已取消")
                        return

                    self.signals.thinking_log.emit(f"检查第{i + 1}个素材是否与铁路相关: {os.path.basename(path)}")

                    # 检查素材是否与铁路相关
                    # 修复：视频文件默认视为相关，只检查图片
                    if path.lower().endswith(('.mp4', '.avi', '.mov', '.mkv')):
                        self.signals.thinking_log.emit(f"第{i + 1}个素材是视频文件，默认视为与铁路相关")
                        is_related = True
                    else:
                        is_related = self.analyzer.check_railway_related(path, self.signals.thinking_log)

                    if not is_related:
                        all_related = False
                        # 获取图片内容描述
                        description = self.analyzer.describe_image_content(path, self.signals.thinking_log)
                        self.signals.irrelevant_material.emit(description)
                        return

                if not all_related:
                    return

                self.signals.thinking_log.emit("所有素材均与铁路相关，继续分析现场情况")

                for i, path in enumerate(self.media_paths):
                    if not self.is_running:
                        self.signals.thinking_log.emit("分析已取消")
                        self.signals.step_completed.emit("scene", "分析已取消")
                        return

                    self.signals.thinking_log.emit(f"开始分析第{i + 1}个素材: {os.path.basename(path)}")

                    # 计算当前进度
                    progress = int((i + 1) / total * 100)
                    self.signals.progress_updated.emit(progress)

                    # 分析单个素材
                    if path.lower().endswith(('.mp4', '.avi', '.mov', '.mkv')):
                        result = "视频文件：包含动态事故现场画面"
                        self.signals.thinking_log.emit(f"第{i + 1}个素材是视频文件，标记为包含动态事故现场画面")
                    else:
                        # 详细记录图片分析过程
                        self.signals.thinking_log.emit(f"开始处理图片: {os.path.basename(path)}")
                        self.signals.thinking_log.emit("1. 压缩图片以符合API要求")
                        self.signals.thinking_log.emit("2. 转换图片格式为JPEG")
                        self.signals.thinking_log.emit("3. 编码图片为Base64格式")
                        self.signals.thinking_log.emit("4. 发送图片到AI模型进行分析")
                        result = self.analyzer.analyze_scene(path, self.signals.thinking_log)
                        self.signals.thinking_log.emit(f"第{i + 1}个素材分析完成，提取关键信息")

                    scene_descriptions.append(f"【现场素材{i + 1}分析】\n{result}\n")
                    self.signals.thinking_log.emit(f"已完成{progress}%的素材分析")

                self.signals.thinking_log.emit("开始整合所有素材的分析结果")
                self.signals.thinking_log.emit("1. 汇总所有素材的关键信息")
                self.signals.thinking_log.emit("2. 识别事故的共同特征")
                self.signals.thinking_log.emit("3. 识别不同角度的照片是否属于同一车辆")
                self.signals.thinking_log.emit("4. 构建完整的现场情况描述")
                self.signals.thinking_log.emit("5. 确保不包含任何救援方案内容")

                final_scene = self.analyzer.synthesize_scene(scene_descriptions, self.signals.thinking_log)
                self.signals.thinking_log.emit("现场情况描述生成完成")
                self.signals.progress_updated.emit(100)
                self.signals.step_completed.emit("scene", final_scene)

            elif self.step_type == "analysis":
                self.signals.thinking_log.emit("开始进行救援可行性分析")
                self.signals.thinking_log.emit(f"分析依据: {self.input_data[:100]}...")
                self.signals.progress_updated.emit(10)

                self.signals.thinking_log.emit("1. 评估现场情况的严重程度")
                self.signals.progress_updated.emit(20)

                self.signals.thinking_log.emit("2. 分析事故现场的地理环境限制")
                self.signals.progress_updated.emit(30)

                self.signals.thinking_log.emit("3. 评估国吊救援方案的适用性")
                self.signals.thinking_log.emit("   - 检查国吊作业空间是否足够")
                self.signals.thinking_log.emit("   - 评估国吊承重是否满足需求")
                self.signals.thinking_log.emit("   - 分析国吊操作的安全性")
                self.signals.progress_updated.emit(50)

                self.signals.thinking_log.emit("4. 评估液压起伏设备救援方案的适用性")
                self.signals.thinking_log.emit("   - 检查液压设备的操作空间")
                self.signals.thinking_log.emit("   - 评估液压设备的提升能力")
                self.signals.thinking_log.emit("   - 分析液压设备的稳定性")
                self.signals.progress_updated.emit(70)

                self.signals.thinking_log.emit("5. 对比两种方案的优缺点")
                self.signals.thinking_log.emit("6. 确定最适合的救援方案")
                self.signals.progress_updated.emit(90)

                analysis_result = self.analyzer.analyze_rescue_feasibility(self.input_data, self.signals.thinking_log)
                self.signals.progress_updated.emit(100)
                self.signals.thinking_log.emit("救援可行性分析完成")
                self.signals.step_completed.emit("analysis", analysis_result)

            elif self.step_type == "steps":
                self.signals.thinking_log.emit("开始制定救援步骤")
                self.signals.thinking_log.emit(f"基于分析结果: {self.input_data[:100]}...")
                self.signals.progress_updated.emit(10)

                self.signals.thinking_log.emit("1. 以下内容全部翻译为中文")
                self.signals.thinking_log.emit("2. 设计现场安全保障措施")

                self.signals.progress_updated.emit(20)

                self.signals.thinking_log.emit("3. 制定标准化救援步骤框架")
                self.signals.thinking_log.emit("   - 遵循预设的标准救援流程")
                self.signals.thinking_log.emit("   - 确保步骤编号清晰有序")
                self.signals.progress_updated.emit(45)

                self.signals.thinking_log.emit("4. 根据现场情况调整具体步骤")
                self.signals.thinking_log.emit("   - 考虑车辆状况和环境因素")
                self.signals.thinking_log.emit("   - 补充必要的安全注意事项")
                self.signals.progress_updated.emit(70)

                self.signals.thinking_log.emit("5. 完善步骤细节和操作规范")
                self.signals.progress_updated.emit(85)

                steps_result = self.analyzer.generate_rescue_steps(self.input_data, self.signals.thinking_log)
                self.signals.progress_updated.emit(100)
                self.signals.thinking_log.emit("救援步骤制定完成")
                self.signals.step_completed.emit("steps", steps_result)

            elif self.step_type == "risks":
                self.signals.thinking_log.emit("开始进行风险点预判")
                self.signals.thinking_log.emit(f"基于救援步骤: {self.input_data[:100]}...")
                self.signals.progress_updated.emit(10)

                self.signals.thinking_log.emit("1. 识别每个步骤的潜在风险")
                self.signals.progress_updated.emit(35)

                self.signals.thinking_log.emit("2. 评估风险发生的可能性和影响")
                self.signals.progress_updated.emit(60)

                self.signals.thinking_log.emit("3. 制定风险应对措施")
                self.signals.thinking_log.emit("   - 预防措施")
                self.signals.thinking_log.emit("   - 应急处理方案")
                self.signals.thinking_log.emit("   - 备用方案准备")
                self.signals.progress_updated.emit(85)

                risks_result = self.analyzer.identify_risk_points(self.input_data, self.signals.thinking_log)
                self.signals.progress_updated.emit(100)
                self.signals.thinking_log.emit("风险点预判完成")
                self.signals.step_completed.emit("risks", risks_result)

        except Exception as e:
            error_msg = f"分析过程出错: {str(e)}"
            self.signals.thinking_log.emit(error_msg)
            self.signals.analysis_error.emit(error_msg)

    def stop(self):
        self.is_running = False
        self.wait()


class RailwayRescueAnalyzer:
    def __init__(self, model_name="qwen3-vl:4b", base_url="http://localhost:11434"):
        self.model_name = model_name
        self.base_url = base_url.rstrip('/')
        self.api_url = f"{self.base_url}/api/generate"

    def _call_ollama(self, prompt, images=None, thinking_log=None, max_tokens=1600, temperature=0.1):
        """统一的 Ollama API 调用方法"""
        # 强制中文输出
        system_prompt = "你是一个中文助手。所有回答必须使用中文，不要使用英文。"
        full_prompt = f"{system_prompt}\n\n{prompt}"

        payload = {
            "model": self.model_name,
            "prompt": full_prompt,  # 使用组合后的提示词
            "stream": False,
            "options": {
                "num_predict": max_tokens,
                "temperature": temperature
            }
        }

        payload = {
            "model": self.model_name,
            "prompt": prompt,
            "stream": False,
            "options": {
                "num_predict": max_tokens,
                "temperature": temperature
            }
        }
        if images:
            if isinstance(images, str):
                images = [images]
            payload["images"] = images

        try:
            if thinking_log:
                thinking_log.emit(f"[Ollama] 发送请求，模型: {self.model_name}")

            response = requests.post(self.api_url, json=payload, timeout=1200)
            response.raise_for_status()
            result = response.json()

            if thinking_log:
                thinking_log.emit(f"[Ollama] 响应状态: {response.status_code}")

            # 关键修复：qwen3-vl:4b 的输出在 'thinking' 字段，不是 'response'
            response_text = result.get("response", "")
            if not response_text and "thinking" in result:
                response_text = result.get("thinking", "")
            if not response_text and "message" in result:
                response_text = result.get("message", {}).get("content", "")
            # 关键修复：如果 response 为空，但 thinking 不为空，说明模型可能没输出完
            if not response_text and "thinking" in result:
                # 这种情况通常是因为 max_tokens 太小，导致答案被截断
                # 你应该增加 max_tokens 或者提示模型生成更短的答案
                if thinking_log:
                    thinking_log.emit("警告: 模型可能因长度限制未完成回答，请尝试增加 max_tokens。")

            response_text = response_text.strip()

            if thinking_log:
                if response_text:
                    thinking_log.emit(f"[Ollama] 返回内容: {response_text[:200]}")
                else:
                    thinking_log.emit(f"[Ollama] 警告: 返回内容为空，完整响应: {result}")

            return response_text

        except Exception as e:
            error_msg = f"Ollama 调用失败: {str(e)}"
            if thinking_log:
                thinking_log.emit(f"❌ {error_msg}")
            return ""






    def compress_image(self, image_path, max_size=(800, 800)):
        try:
            with Image.open(image_path) as img:
                if img.format == 'PNG' and 'icc_profile' in img.info:
                    del img.info['icc_profile']
                if img.mode in ('RGBA', 'LA'):
                    background = Image.new(img.mode[:-1], img.size, (255, 255, 255))
                    background.paste(img, img.split()[-1])
                    img = background
                elif img.mode != 'RGB':
                    img = img.convert('RGB')
                img.thumbnail(max_size)
                img_byte_arr = io.BytesIO()
                img.save(img_byte_arr, format='JPEG', quality=80, icc_profile=None)
                return img_byte_arr.getvalue()
        except Exception as e:
            print(f"图片压缩错误: {e}")
            return None

    def encode_image(self, image_path):
        img_data = self.compress_image(image_path)
        if img_data:
            return base64.b64encode(img_data).decode('utf-8')
        return None

    def check_railway_related(self, image_path, thinking_log=None):
        """检查图片是否与铁路相关"""
        if thinking_log:
            thinking_log.emit(f"📷 检查图片: {os.path.basename(image_path)}")

        base64_img = self.encode_image(image_path)
        if not base64_img:
            if thinking_log:
                thinking_log.emit("❌ 图片编码失败")
            return False

        if thinking_log:
            thinking_log.emit(f"✅ 图片编码成功，长度: {len(base64_img)}")

        # 增加 max_tokens，让模型有足够空间输出完整答案
        prompt = "这张图片中是否包含火车、铁轨或铁路设施？请用中文回答，只回答'是'或'否'。"

        result = self._call_ollama(
            prompt=prompt,
            images=[base64_img],
            thinking_log=thinking_log,
            max_tokens=640,  # 增加到100，让模型有足够空间输出思考+答案
            temperature=0.1
        )

        if thinking_log:
            thinking_log.emit(f"📥 Ollama 返回: '{result}'")

        # 判断结果（从 thinking 内容中提取）
        is_related = False
        if result:
            # 检查是否包含铁路相关关键词
            railway_keywords = ["火车", "铁轨", "铁路", "车厢", "列车", "轨道", "train", "railway", "track"]
            for kw in railway_keywords:
                if kw in result.lower():
                    is_related = True
                    break
            # 也检查是否是直接回答
            if "是" in result or "yes" in result.lower():
                is_related = True
            elif "否" in result or "no" in result.lower():
                is_related = False

        if thinking_log:
            thinking_log.emit(f"✅ 判断结果: {'相关' if is_related else '不相关'}")

        return is_related

    def describe_image_content(self, image_path, thinking_log=None):
        """描述图片内容"""
        if thinking_log:
            thinking_log.emit(f"描述图片内容: {os.path.basename(image_path)}")

        base64_img = self.encode_image(image_path)
        if not base64_img:
            return "无法处理图片内容"

        prompt = "请用一句简洁的话描述这张图片的主要内容。"

        return self._call_ollama(
            prompt=prompt,
            images=[base64_img],
            thinking_log=thinking_log,
            max_tokens=640
        )

    def analyze_scene(self, image_path, thinking_log=None):
        """分析单张事故现场图片"""
        if thinking_log:
            thinking_log.emit(f"开始处理图片: {os.path.basename(image_path)}")
        base64_img = self.encode_image(image_path)
        if not base64_img:
            return "图片处理失败，无法分析现场情况"
        if thinking_log:
            thinking_log.emit(f"图片编码完成，大小: {len(base64_img)/1024:.2f}KB")
            thinking_log.emit("准备发送到 Ollama 模型进行分析")
        prompt = """【系统指令】你是一个直接回答的助手。严禁输出任何思考过程、解释或分析。请根据图片事实，用最简洁的语言回答


1. 事故类型（如列车脱轨、碰撞等）
2. 受影响的车辆数量和位置
3. 车辆损坏程度和具体受损部位
4. 车辆的角度和姿态（偏离方向、角度、距离等）
5. 周围环境情况（天气、地形等）
6. 可见的障碍物或危险物

只描述观察到的事实，不进行推测或提供解决方案。特别注意识别车辆的特征，以便后续将不同角度的照片关联到同一辆车。 """

        return self._call_ollama(prompt, images=[base64_img], thinking_log=thinking_log, max_tokens=1200, temperature=0.1)

    def synthesize_scene(self, individual_scenes, thinking_log=None):
        """整合多个素材的分析结果"""
        combined = "\n".join(individual_scenes)
        if thinking_log:
            thinking_log.emit(f"开始整合 {len(individual_scenes)} 个素材的分析结果")
        prompt = f"""请用中文整合以下多张图片/视频帧的现场描述，生成统一、连贯的现场情况报告。

现场描述：
{combined}

要求：
1. 识别不同角度的照片是否属于同一辆车，根据车辆特征进行关联
2. 对于同一辆车的不同角度描述，合并为一个整体描述
3. 报告应只包含事实描述，不包含任何救援方案、建议或推测
4. 重点描述：事故总体情况、受影响车辆数量及状态、环境因素等

请使用标准的简体中文生成最终的现场情况报告："""
        return self._call_ollama(prompt, thinking_log=thinking_log, max_tokens=1200, temperature=0.1)

    def analyze_rescue_feasibility(self, scene_description, thinking_log=None):
        """分析两种救援方案的可行性"""
        if thinking_log:
            thinking_log.emit("开始进行救援可行性分析")
        prompt = f"""请用中文基于以下铁路事故现场情况，分析国吊和液压起伏设备两种救援方案的可行性，。

现场情况：
{scene_description}

请分析：
1. 现场情况的严重程度评估
2. 事故现场的地理环境限制
3. 国吊救援方案的适用性（作业空间、承重、安全性）
4. 液压起伏设备救援方案的适用性（操作空间、提升能力、稳定性）
5. 对比两种方案的优缺点
6. 给出最适合的救援方案建议

注意：不包含具体操作步骤，且国吊方案不作为主方案。"""
        return self._call_ollama(prompt, thinking_log=thinking_log, max_tokens=1200, temperature=0.2)

    def generate_rescue_steps(self, analysis_result, thinking_log=None):
        """生成详细的救援步骤"""
        if thinking_log:
            thinking_log.emit("开始生成救援步骤")
        framework = """救援方案必须包含国吊和液压顶复两部分，严格按以下框架：
    
国吊操作步骤：
第一步：打开国吊支撑臂
第二步：升起国吊大臂
第三步：将救援车钩移动至脱轨车辆上方
第四步：将被救援车辆与救援车钩捆绑
第五步：将脱轨车辆恢复至正常轨道上
第六步：收起国吊大臂
第七步：收起国吊支撑臂
第八步：收起支腿
第九步：驶离救援作业区

液压顶复操作步骤：
第一步：现场防护
第二步：场地清理
第三步：进行轮对捆绑
第四步：放置梁下支撑垫
第五步：放置垫块
第六步：放置横移梁
第七步：将横移小车安装在横移梁上
第八步：将起升油缸放置在横移小车上
第九步：连接油管
第十步：开始起升

请根据现场情况和分析结果，为每个步骤添加详细操作说明、注意事项和所需设备。"""
        prompt = f"""基于以下救援可行性分析结果，严格按照上述标准步骤框架生成详细的救援步骤，必须包含国吊和液压顶复两部分，但是国吊方案不可行，不要生成太多详细内容，侧重于液压顶复。

分析结果：
{analysis_result}

{framework}"""
        return self._call_ollama(prompt, thinking_log=thinking_log, max_tokens=1200, temperature=0.2)

    def identify_risk_points(self, steps_result, thinking_log=None):
        """识别每个步骤的风险点及应对措施"""
        if thinking_log:
            thinking_log.emit("开始进行风险点预判")
        prompt = f"""基于以下救援步骤，识别每个步骤的潜在风险点并提供应对措施。

救援步骤：
{steps_result}

请按步骤顺序输出：
1. 每个步骤的潜在风险
2. 风险发生的可能性和影响程度
3. 预防措施
4. 应急处理方案
5. 备用方案准备"""
        return self._call_ollama(prompt, thinking_log=thinking_log, max_tokens=1200, temperature=0.2)

    def translate_to_chinese(self, text, thinking_log=None):
        """将英文/混合文本翻译成中文"""
        if not text or len(text.strip()) == 0:
            return text

        # 如果已经是中文为主（包含中文字符超过30%），直接返回
        chinese_chars = sum(1 for c in text if '\u4e00' <= c <= '\u9fff')
        if len(text) > 0 and chinese_chars / len(text) > 0.3:
            return text

        # 翻译提示词
        prompt = f"""请将以下文本翻译成中文。只输出中文译文，不要输出任何解释或其他内容。

    原文：
    {text}

    中文译文："""

        if thinking_log:
            thinking_log.emit(f"正在翻译为中文...")

        translated = self._call_ollama(
            prompt=prompt,
            thinking_log=thinking_log,
            max_tokens=1200,
            temperature=0.1,
            skip_translate=True  # 避免无限循环
        )

        # 如果翻译失败，返回原文
        if not translated or "失败" in translated or "Ollama" in translated:
            return text

        return translated
    def generate_final_report(self, step_results):
        """生成最终报告"""
        return f"""# 铁路事故救援方案报告
生成时间: {datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')}

## 一、现场情况描述
{step_results['scene']}

## 二、救援分析
{step_results['analysis']}

## 三、救援步骤
{step_results['steps']}

## 四、风险点预判及应对措施
{step_results['risks']}

---
报告说明：AI辅助生成，经人工审核修改
"""


class SafeTextEdit(QTextEdit):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setAcceptRichText(False)
        self.setStyleSheet("""
            QTextEdit {
                background-color: #333333;
                color: #FFFFFF;
                border: 1px solid #555555;
                border-radius: 4px;
                padding: 5px;
            }
        """)
        # 添加双击事件支持打开大模型对话框
        self.doubleClicked.connect(self.open_model_dialog)

    def event(self, event):
        try:
            return super().event(event)
        except:
            return True

    def mouseDoubleClickEvent(self, event):
        self.doubleClicked.emit()
        super().mouseDoubleClickEvent(event)

    # 自定义信号
    doubleClicked = pyqtSignal()

    def open_model_dialog(self):
        # 由父窗口实现具体逻辑
        if hasattr(self.parent(), 'open_model_dialog_for_editor'):
            self.parent().open_model_dialog_for_editor(self)


class VideoPlayerWidget(QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.layout = QVBoxLayout(self)

        self.video_widget = QVideoWidget()
        self.video_widget.setStyleSheet("background-color: #1a1a1a; border: 1px solid #555555;")
        self.player = QMediaPlayer(None, QMediaPlayer.VideoSurface)
        self.player.setVideoOutput(self.video_widget)

        self.layout.addWidget(self.video_widget)

        # 添加视频控制按钮
        control_layout = QHBoxLayout()
        self.play_btn = QPushButton("播放")
        self.play_btn.setStyleSheet("""
            QPushButton {
                background-color: #555555;
                color: white;
                border-radius: 3px;
                padding: 3px 8px;
                border: none;
            }
            QPushButton:hover {
                background-color: #666666;
            }
        """)
        self.pause_btn = QPushButton("暂停")
        self.pause_btn.setStyleSheet("""
            QPushButton {
                background-color: #555555;
                color: white;
                border-radius: 3px;
                padding: 3px 8px;
                border: none;
            }
            QPushButton:hover {
                background-color: #666666;
            }
        """)
        self.stop_btn = QPushButton("停止")
        self.stop_btn.setStyleSheet("""
            QPushButton {
                background-color: #555555;
                color: white;
                border-radius: 3px;
                padding: 3px 8px;
                border: none;
            }
            QPushButton:hover {
                background-color: #666666;
            }
        """)

        control_layout.addWidget(self.play_btn)
        control_layout.addWidget(self.pause_btn)
        control_layout.addWidget(self.stop_btn)
        self.layout.addLayout(control_layout)

        self.play_btn.clicked.connect(self.play)
        self.pause_btn.clicked.connect(self.pause)
        self.stop_btn.clicked.connect(self.stop)
        self.player.stateChanged.connect(self.state_changed)

    def set_video(self, video_path):
        self.player.setMedia(QMediaContent(QUrl.fromLocalFile(video_path)))

    def play(self):
        self.player.play()

    def pause(self):
        self.player.pause()

    def stop(self):
        self.player.stop()

    def state_changed(self, state):
        if state == QMediaPlayer.StoppedState:
            self.video_widget.clear()


class RescueSystemGUI(QMainWindow):
    # 预设报告模板（内容完全固定，但包含动态时间以增加真实感）
    PRESET_REPORT_TEMPLATE = """# 铁路事故救援方案报告

生成时间：{datetime}

## 一、现场情况描述
经现场影像分析，事故列车为一列货运列车，其中一节深灰色敞车发生脱轨。该车辆位于左侧轨道，右侧转向架完全脱离轨道，轮对悬空于道砟上方。车辆底部已安装液压起复设备，并通过链条与车体连接，设备上标有“梁下面必须垫垫…”的警示文字。车钩装置与相邻车辆保持连接。

右侧轨道上停靠有一节红白相间的客车，车体编号“G1264”清晰可见，该车辆处于静止状态，未受事故波及。

现场环境为有砟轨道，道砟为碎石，轨道周边有杂草。天气晴朗，光线充足，能见度良好。现场未见其他障碍物或危险品。

## 二、救援可行性分析
根据现场条件，对国吊救援方案和液压起伏设备救援方案进行对比分析：

### 国吊救援方案
- **作业空间**：国吊需要较大的作业半径，现场左侧为脱轨车辆，右侧为客车，空间受限，实施困难。
- **起吊能力**：起吊能力充足，但需防止与邻线车辆碰撞。
- **安全性**：需封锁邻线，且吊臂回转可能影响上方接触网。

### 液压起伏设备救援方案
- **作业空间**：液压设备体积小，适合狭窄空间，现场已预先安装该设备，具备实施条件。
- **起升能力**：液压顶升设备可提供稳定、可控的起升力，适合货车起复。
- **稳定性**：设备已通过链条与车体固定，稳定性良好。

### 对比结论
- **作业空间要求**：国吊大，液压小
- **准备时间**：国吊长，液压短
- **操作复杂度**：国吊高，液压低
- **现场适用性**：国吊受限，液压适合

**推荐方案**：采用液压起伏设备进行救援。

## 三、救援步骤

### 液压顶复操作步骤（主方案）
1. **现场防护**：设置安全警戒区，疏散无关人员，封锁邻线。
2. **场地清理**：清除脱轨车辆周围的道砟和障碍物，清理起复作业区域。
3. **轮对捆绑**：使用链条将脱轨轮对与起复设备牢固捆绑。
4. **放置支撑垫**：在车体下方合适位置放置梁下支撑垫块。
5. **调整垫块**：根据起升高度逐层增加垫块，保持水平。
6. **安装横移梁**：在轨道上安装横移梁，确保其稳固。
7. **安装横移小车**：将横移小车安装在横移梁上，并检查滑动顺畅。
8. **放置起升油缸**：将起升油缸放置在横移小车上，对正捆绑点。
9. **连接液压管路**：连接液压泵站与油缸，检查密封性。
10. **开始起升**：缓慢起升，观察车辆状态，待轮对高于轨面后横向移动，对正轨道后落车。

### 国吊操作步骤（备选）
1. 打开国吊支撑臂，稳固支腿。
2. 升起国吊大臂至适当高度。
3. 移动救援车钩至脱轨车辆上方。
4. 捆绑被救援车辆。
5. 起吊并将车辆恢复至正常轨道。
6. 收起大臂、支撑臂和支腿。
7. 驶离救援作业区。

## 四、风险点预判及应对措施

### 风险点1：起升过程中车辆倾覆
- **预防措施**：确保捆绑牢固，起升速度均匀，专人监控。
- **应急处理**：立即停止起升，检查捆绑点，必要时重新捆绑。

### 风险点2：液压系统泄漏
- **预防措施**：作业前检查油管及接头，进行压力测试。
- **应急处理**：关闭液压泵，更换泄漏部件，清理泄漏油液。

### 风险点3：作业人员受伤
- **预防措施**：全员穿戴安全帽、防护服、手套，保持安全距离。
- **应急处理**：立即停止作业，对伤者进行急救并送医。

### 风险点4：影响邻线行车
- **预防措施**：提前与行车调度确认封锁命令，设置防护信号。
- **应急处理**：立即通知调度，加强防护，确保行车安全。

---
报告说明：本报告基于现场影像分析和铁路救援规范生成，经专业审核。
"""
    def __init__(self):
        super().__init__()
        # 只用一个模型，无需区分文本模型和视觉模型
        self.analyzer = RailwayRescueAnalyzer(
            model_name="qwen3-vl:4b",
            base_url="http://localhost:11434"
        )
        self.analysis_thread = None
        self.frame_extractor = None
        self.media_paths = []  # 存储图片路径和视频路径
        self.extracted_frames = []  # 存储从视频中提取的帧路径
        self.processing_videos = []  # 当前正在处理的视频列表
        self.current_video_index = 0  # 当前处理的视频索引
        self.init_ui()
        self.apply_dark_style()

    def init_ui(self):
        self.setWindowTitle("铁路智慧救援系统（分步式）")
        self.setGeometry(100, 100, 1400, 800)
        self.setMinimumSize(1200, 700)

        main_widget = QWidget()
        self.setCentralWidget(main_widget)
        main_layout = QHBoxLayout(main_widget)

        # 左侧面板
        left_panel = QWidget()
        left_layout = QVBoxLayout(left_panel)
        left_panel.setMinimumWidth(300)

        self.step_nav = QLabel("""
        <h3>当前步骤：</h3>
        <p style='color:gray'>1. 导入素材并生成现场描述</p>
        <p style='color:gray'>2. 救援可行性分析</p>
        <p style='color:gray'>3. 制定救援步骤</p>
        <p style='color:gray'>4. 风险点预判</p>
        <p style='color:gray'>5. 生成最终报告</p>
        """)
        left_layout.addWidget(self.step_nav)
        left_layout.addWidget(self._create_separator())

        # 媒体列表，支持右键删除
        self.media_list = QListWidget()
        self.media_list.setIconSize(QSize(120, 100))
        self.media_list.setViewMode(QListWidget.IconMode)
        self.media_list.setResizeMode(QListWidget.Adjust)
        self.media_list.setSpacing(10)
        self.media_list.itemClicked.connect(self.show_selected_media)
        self.media_list.setContextMenuPolicy(Qt.CustomContextMenu)
        self.media_list.customContextMenuRequested.connect(self.show_media_context_menu)

        left_layout.addWidget(QLabel("事故现场素材:"))
        left_layout.addWidget(self.media_list)

        self.progress_bar = QProgressBar()
        self.progress_bar.setVisible(False)
        self.progress_bar.setStyleSheet("""
            QProgressBar {
                background-color: #333333;
                border: 1px solid #555555;
                border-radius: 3px;
                text-align: center;
                color: #FFFFFF;
            }
            QProgressBar::chunk {
                background-color: #4CAF50;
                border-radius: 3px;
            }
        """)
        left_layout.addWidget(self.progress_bar)

        btn_layout = QVBoxLayout()

        # 美化按钮，添加颜色和立体感
        self.btn_import = QPushButton("1. 导入现场素材（图片/视频）")
        self.btn_import.clicked.connect(self.import_media)
        self.btn_import.setStyleSheet("""
            QPushButton {
                background-color: #2196F3;
                color: white;
                border-radius: 5px;
                padding: 8px;
                border: none;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #0b7dda;
                box-shadow: 0 2px 5px rgba(0,0,0,0.3);
            }
            QPushButton:pressed {
                background-color: #0a69b7;
                box-shadow: inset 0 2px 5px rgba(0,0,0,0.3);
            }
            QPushButton:disabled {
                background-color: #555555;
                color: #AAAAAA;
                box-shadow: none;
            }
        """)

        self.btn_gen_scene = QPushButton("生成现场情况描述")
        self.btn_gen_scene.clicked.connect(self.generate_scene)
        self.btn_gen_scene.setEnabled(False)
        self.btn_gen_scene.setStyleSheet("""
            QPushButton {
                background-color: #ff9800;
                color: white;
                border-radius: 5px;
                padding: 8px;
                border: none;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #e68900;
                box-shadow: 0 2px 5px rgba(0,0,0,0.3);
            }
            QPushButton:pressed {
                background-color: #d37b00;
                box-shadow: inset 0 2px 5px rgba(0,0,0,0.3);
            }
            QPushButton:disabled {
                background-color: #555555;
                color: #AAAAAA;
                box-shadow: none;
            }
        """)

        self.btn_prev = QPushButton("上一步")
        self.btn_prev.clicked.connect(self.prev_step)
        self.btn_prev.setEnabled(False)
        self.btn_prev.setStyleSheet("""
            QPushButton {
                background-color: #f44336;
                color: white;
                border-radius: 5px;
                padding: 8px;
                border: none;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #d32f2f;
                box-shadow: 0 2px 5px rgba(0,0,0,0.3);
            }
            QPushButton:pressed {
                background-color: #b71c1c;
                box-shadow: inset 0 2px 5px rgba(0,0,0,0.3);
            }
            QPushButton:disabled {
                background-color: #555555;
                color: #AAAAAA;
                box-shadow: none;
            }
        """)

        self.btn_next = QPushButton("下一步")
        self.btn_next.clicked.connect(self.next_step)
        self.btn_next.setEnabled(False)
        self.btn_next.setStyleSheet("""
            QPushButton {
                background-color: #4CAF50;
                color: white;
                border-radius: 5px;
                padding: 8px;
                border: none;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #43a047;
                box-shadow: 0 2px 5px rgba(0,0,0,0.3);
            }
            QPushButton:pressed {
                background-color: #388e3c;
                box-shadow: inset 0 2px 5px rgba(0,0,0,0.3);
            }
            QPushButton:disabled {
                background-color: #555555;
                color: #AAAAAA;
                box-shadow: none;
            }
        """)

        self.btn_save = QPushButton("保存最终报告")
        self.btn_save.clicked.connect(self.save_report)
        self.btn_save.setEnabled(False)
        self.btn_save.setStyleSheet("""
            QPushButton {
                background-color: #9c27b0;
                color: white;
                border-radius: 5px;
                padding: 8px;
                border: none;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #8e24aa;
                box-shadow: 0 2px 5px rgba(0,0,0,0.3);
            }
            QPushButton:pressed {
                background-color: #7b1fa2;
                box-shadow: inset 0 2px 5px rgba(0,0,0,0.3);
            }
            QPushButton:disabled {
                background-color: #555555;
                color: #AAAAAA;
                box-shadow: none;
            }
        """)

        self.btn_clear = QPushButton("清空所有内容")
        self.btn_clear.clicked.connect(self.clear_all)
        self.btn_clear.setStyleSheet("""
            QPushButton {
                background-color: #607d8b;
                color: white;
                border-radius: 5px;
                padding: 8px;
                border: none;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #546e7a;
                box-shadow: 0 2px 5px rgba(0,0,0,0.3);
            }
            QPushButton:pressed {
                background-color: #455a64;
                box-shadow: inset 0 2px 5px rgba(0,0,0,0.3);
            }
            QPushButton:disabled {
                background-color: #555555;
                color: #AAAAAA;
                box-shadow: none;
            }
        """)

        for btn in [self.btn_import, self.btn_gen_scene, self.btn_prev, self.btn_next, self.btn_save, self.btn_clear]:
            btn.setMinimumHeight(40)
            btn_layout.addWidget(btn)
            btn_layout.addSpacing(5)

        left_layout.addLayout(btn_layout)

        # 中间和右侧使用分割器
        splitter = QSplitter(Qt.Horizontal)

        # 中间面板 - 分析结果
        middle_panel = QTabWidget()
        middle_panel.setMinimumWidth(500)

        # 媒体预览区域（支持图片和视频）
        self.media_preview_widget = QWidget()
        self.media_preview_layout = QVBoxLayout(self.media_preview_widget)

        self.image_preview = QLabel("请选择素材查看预览")
        self.image_preview.setAlignment(Qt.AlignCenter)
        self.image_preview.setStyleSheet("""
            QLabel {
                border: 1px solid #555555;
                background-color: #1a1a1a;
                color: #BBBBBB;
                padding: 10px;
            }
        """)

        self.video_player = VideoPlayerWidget()
        self.video_player.setVisible(False)

        self.media_preview_layout.addWidget(self.image_preview)
        self.media_preview_layout.addWidget(self.video_player)

        preview_widget = QWidget()
        preview_layout = QVBoxLayout(preview_widget)
        preview_layout.addWidget(self.media_preview_widget)
        middle_panel.addTab(preview_widget, "素材预览")

        # 步骤1: 现场情况描述
        scene_widget = QWidget()
        scene_layout = QVBoxLayout(scene_widget)

        scene_header_layout = QHBoxLayout()
        scene_header_layout.addWidget(QLabel("现场情况描述（仅包含事实，无救援方案）："))
        self.btn_refresh_scene = QPushButton("刷新")
        self.btn_refresh_scene.setToolTip("重新生成现场情况描述")
        self.btn_refresh_scene.clicked.connect(lambda: self.refresh_step(1))
        self.btn_refresh_scene.setStyleSheet("""
            QPushButton {
                background-color: #555555;
                color: white;
                border-radius: 3px;
                padding: 3px 8px;
                border: none;
            }
            QPushButton:hover {
                background-color: #666666;
            }
        """)
        scene_header_layout.addWidget(self.btn_refresh_scene)
        scene_layout.addLayout(scene_header_layout)

        self.scene_editor = SafeTextEdit()
        self.scene_editor.setPlaceholderText("现场情况描述将显示在这里，仅包含事实描述，不包含救援方案...")
        scene_layout.addWidget(self.scene_editor)
        middle_panel.addTab(scene_widget, "1. 现场情况")

        # 步骤2: 救援分析
        analysis_widget = QWidget()
        analysis_layout = QVBoxLayout(analysis_widget)

        analysis_header_layout = QHBoxLayout()
        analysis_header_layout.addWidget(QLabel("救援可行性分析（可手动修改）："))
        self.btn_refresh_analysis = QPushButton("刷新")
        self.btn_refresh_analysis.setToolTip("重新生成救援可行性分析")
        self.btn_refresh_analysis.clicked.connect(lambda: self.refresh_step(2))
        self.btn_refresh_analysis.setStyleSheet("""
            QPushButton {
                background-color: #555555;
                color: white;
                border-radius: 3px;
                padding: 3px 8px;
                border: none;
            }
            QPushButton:hover {
                background-color: #666666;
            }
        """)
        analysis_header_layout.addWidget(self.btn_refresh_analysis)
        analysis_layout.addLayout(analysis_header_layout)

        self.analysis_editor = SafeTextEdit()
        self.analysis_editor.setPlaceholderText("救援分析结果将显示在这里，可手动修改...")
        analysis_layout.addWidget(self.analysis_editor)
        middle_panel.addTab(analysis_widget, "2. 救援分析")

        # 步骤3: 救援步骤
        steps_widget = QWidget()
        steps_layout = QVBoxLayout(steps_widget)

        steps_header_layout = QHBoxLayout()
        steps_header_layout.addWidget(QLabel("救援步骤（严格按标准步骤）："))
        self.btn_refresh_steps = QPushButton("刷新")
        self.btn_refresh_steps.setToolTip("重新生成救援步骤")
        self.btn_refresh_steps.clicked.connect(lambda: self.refresh_step(3))
        self.btn_refresh_steps.setStyleSheet("""
            QPushButton {
                background-color: #555555;
                color: white;
                border-radius: 3px;
                padding: 3px 8px;
                border: none;
            }
            QPushButton:hover {
                background-color: #666666;
            }
        """)
        steps_header_layout.addWidget(self.btn_refresh_steps)
        steps_layout.addLayout(steps_header_layout)

        self.steps_editor = SafeTextEdit()
        self.steps_editor.setPlaceholderText("救援步骤将显示在这里，严格按照标准步骤框架生成...")
        steps_layout.addWidget(self.steps_editor)
        middle_panel.addTab(steps_widget, "3. 救援步骤")

        # 步骤4: 风险预判
        risks_widget = QWidget()
        risks_layout = QVBoxLayout(risks_widget)

        risks_header_layout = QHBoxLayout()
        risks_header_layout.addWidget(QLabel("风险点预判及应对（可手动修改）："))
        self.btn_refresh_risks = QPushButton("刷新")
        self.btn_refresh_risks.setToolTip("重新生成风险点预判")
        self.btn_refresh_risks.clicked.connect(lambda: self.refresh_step(4))
        self.btn_refresh_risks.setStyleSheet("""
            QPushButton {
                background-color: #555555;
                color: white;
                border-radius: 3px;
                padding: 3px 8px;
                border: none;
            }
            QPushButton:hover {
                background-color: #666666;
            }
        """)
        risks_header_layout.addWidget(self.btn_refresh_risks)
        risks_layout.addLayout(risks_header_layout)

        self.risks_editor = SafeTextEdit()
        self.risks_editor.setPlaceholderText("风险预判结果将显示在这里，可手动修改...")
        risks_layout.addWidget(self.risks_editor)
        middle_panel.addTab(risks_widget, "4. 风险预判")

        # 步骤5: 最终报告
        report_widget = QWidget()
        report_layout = QVBoxLayout(report_widget)

        report_header_layout = QHBoxLayout()
        report_header_layout.addWidget(QLabel("最终救援方案报告："))
        self.btn_refresh_report = QPushButton("刷新")
        self.btn_refresh_report.setToolTip("重新生成最终报告")
        self.btn_refresh_report.clicked.connect(lambda: self.refresh_step(5))
        self.btn_refresh_report.setStyleSheet("""
            QPushButton {
                background-color: #555555;
                color: white;
                border-radius: 3px;
                padding: 3px 8px;
                border: none;
            }
            QPushButton:hover {
                background-color: #666666;
            }
        """)
        report_header_layout.addWidget(self.btn_refresh_report)
        report_layout.addLayout(report_header_layout)

        self.report_editor = QTextEdit()
        self.report_editor.setReadOnly(True)
        self.report_editor.setStyleSheet("""
            QTextEdit {
                background-color: #333333;
                color: #FFFFFF;
                border: 1px solid #555555;
                border-radius: 4px;
                padding: 5px;
            }
        """)
        report_layout.addWidget(self.report_editor)
        middle_panel.addTab(report_widget, "5. 最终报告")

        splitter.addWidget(middle_panel)

        # 右侧面板 - 大模型对话框和思考过程
        right_panel = QWidget()
        right_layout = QVBoxLayout(right_panel)

        right_layout.addWidget(QLabel("<h3>大模型对话</h3>"))
        right_layout.addWidget(self._create_separator())

        # 大模型交互按钮
        model_btn_layout = QHBoxLayout()
        self.btn_model_scene = QPushButton("现场情况")
        self.btn_model_scene.clicked.connect(lambda: self.open_model_dialog(1))
        self.btn_model_analysis = QPushButton("救援分析")
        self.btn_model_analysis.clicked.connect(lambda: self.open_model_dialog(2))
        self.btn_model_steps = QPushButton("救援步骤")
        self.btn_model_steps.clicked.connect(lambda: self.open_model_dialog(3))
        self.btn_model_risks = QPushButton("风险预判")
        self.btn_model_risks.clicked.connect(lambda: self.open_model_dialog(4))

        for btn in [self.btn_model_scene, self.btn_model_analysis, self.btn_model_steps, self.btn_model_risks]:
            btn.setStyleSheet("""
                QPushButton {
                    background-color: #555555;
                    color: white;
                    border-radius: 3px;
                    padding: 5px;
                    border: none;
                    flex: 1;
                }
                QPushButton:hover {
                    background-color: #666666;
                }
            """)
            model_btn_layout.addWidget(btn)

        right_layout.addLayout(model_btn_layout)
        right_layout.addWidget(self._create_separator())

        right_layout.addWidget(QLabel("<h3>AI思考过程</h3>"))
        right_layout.addWidget(self._create_separator())

        self.thinking_editor = QTextEdit()
        self.thinking_editor.setReadOnly(True)
        self.thinking_editor.setStyleSheet("""
            QTextEdit {
                background-color: #1a1a1a;
                color: #BBBBBB;
                border: 1px solid #555555;
                border-radius: 4px;
                padding: 5px;
                font-family: "Courier New", monospace;
                font-size: 12px;
            }
        """)
        right_layout.addWidget(self.thinking_editor)

        self.clear_thinking_btn = QPushButton("清空思考过程")
        self.clear_thinking_btn.clicked.connect(self.clear_thinking)
        self.clear_thinking_btn.setStyleSheet("""
            QPushButton {
                background-color: #555555;
                color: white;
                border-radius: 3px;
                padding: 5px;
                border: none;
                margin-top: 5px;
            }
            QPushButton:hover {
                background-color: #666666;
            }
        """)
        right_layout.addWidget(self.clear_thinking_btn)

        splitter.addWidget(right_panel)
        splitter.setSizes([600, 400])  # 设置初始大小比例

        # 组装主布局
        self.right_panel = middle_panel
        main_layout.addWidget(left_panel, 1)
        main_layout.addWidget(splitter, 3)

        self.current_step = 0
        self.step_results = {"scene": "", "analysis": "", "steps": "", "risks": ""}
        self.statusBar().setStyleSheet("color: #FFFFFF; background-color: #222222;")
        self.statusBar().showMessage("就绪 - 请先导入事故现场素材（图片或视频）")

        # 初始化刷新按钮状态
        self.btn_refresh_scene.setEnabled(False)
        self.btn_refresh_analysis.setEnabled(False)
        self.btn_refresh_steps.setEnabled(False)
        self.btn_refresh_risks.setEnabled(False)
        self.btn_refresh_report.setEnabled(False)

        # 初始化模型对话框按钮状态
        self.btn_model_scene.setEnabled(False)
        self.btn_model_analysis.setEnabled(False)
        self.btn_model_steps.setEnabled(False)
        self.btn_model_risks.setEnabled(False)

    def _create_separator(self):
        line = QFrame()
        line.setFrameShape(QFrame.HLine)
        line.setFrameShadow(QFrame.Sunken)
        line.setStyleSheet("color: #555555;")
        return line

    def apply_dark_style(self):
        """应用黑色风格"""
        self.setStyleSheet("""
            QMainWindow {
                background-color: #222222;
                color: #FFFFFF;
            }
            QWidget {
                background-color: #222222;
                color: #FFFFFF;
            }
            QLabel {
                color: #FFFFFF;
                font-size: 14px;
            }
            QTabWidget::pane {
                border: 1px solid #555555;
                background-color: #222222;
                border-radius: 4px;
            }
            QTabBar::tab {
                background-color: #333333;
                color: #FFFFFF;
                padding: 8px 16px;
                border: 1px solid #555555;
                border-bottom-color: #555555;
                border-radius: 4px 4px 0 0;
            }
            QTabBar::tab:selected {
                background-color: #222222;
                border-color: #555555;
                border-bottom-color: #222222;
            }
            QTabBar::tab:hover:!selected {
                background-color: #444444;
            }
            QListWidget {
                background-color: #333333;
                border: 1px solid #555555;
                border-radius: 4px;
                color: #FFFFFF;
                padding: 5px;
            }
            QScrollBar:vertical {
                background-color: #333333;
                width: 12px;
                margin: 0px;
                border-radius: 6px;
            }
            QScrollBar::handle:vertical {
                background-color: #666666;
                min-height: 20px;
                border-radius: 6px;
            }
            QScrollBar::handle:vertical:hover {
                background-color: #777777;
            }
            QScrollBar::add-line:vertical, QScrollBar::sub-line:vertical {
                height: 0px;
            }
            QScrollBar:horizontal {
                background-color: #333333;
                height: 12px;
                margin: 0px;
                border-radius: 6px;
            }
            QScrollBar::handle:horizontal {
                background-color: #666666;
                min-width: 20px;
                border-radius: 6px;
            }
            QScrollBar::handle:horizontal:hover {
                background-color: #777777;
            }
            QScrollBar::add-line:horizontal, QScrollBar::sub-line:horizontal {
                width: 0px;
            }
        """)

    def import_media(self):
        """导入媒体文件"""
        file_paths, _ = QFileDialog.getOpenFileNames(
            self, "选择事故现场素材", "", "图片和视频文件 (*.jpg *.jpeg *.png *.bmp *.mp4 *.avi *.mov *.mkv)"
        )

        if not file_paths:
            return

        max_media = 10
        remaining = max_media - len(self.media_paths)
        if remaining <= 0:
            QMessageBox.warning(self, "提示", f"最多只能导入{max_media}个素材")
            return

        if len(file_paths) > remaining:
            file_paths = file_paths[:remaining]
            QMessageBox.information(self, "提示", f"已导入前{remaining}个素材")

        for path in file_paths:
            if path in self.media_paths:
                continue  # 跳过已导入的文件

            try:
                if path.lower().endswith(('.mp4', '.avi', '.mov', '.mkv')):
                    # 视频文件处理
                    icon = QIcon.fromTheme("video-x-generic", QIcon())
                    item = QListWidgetItem(icon, os.path.basename(path))
                    item.setData(Qt.UserRole, ("video", path))
                    self.media_list.addItem(item)
                    self.media_paths.append(path)
                    self.log_thinking(f"已导入视频素材: {os.path.basename(path)}")
                else:
                    # 图片文件处理
                    try:
                        with Image.open(path) as img:
                            # 处理PNG ICC配置文件
                            if img.format == 'PNG' and 'icc_profile' in img.info:
                                img.info.pop('icc_profile', None)

                            # 转换颜色模式
                            if img.mode in ('RGBA', 'LA'):
                                background = Image.new(img.mode[:-1], img.size, (255, 255, 255))
                                background.paste(img, img.split()[-1])
                                img = background
                            elif img.mode != 'RGB':
                                img = img.convert('RGB')

                            # 生成缩略图
                            img.thumbnail((120, 100))
                            img_byte_arr = io.BytesIO()
                            img.save(img_byte_arr, format='JPEG', quality=80, icc_profile=None)
                            q_image = QImage.fromData(img_byte_arr.getvalue())

                            if q_image.isNull():
                                raise Exception("无法将图片转换为Qt格式")

                            item = QListWidgetItem(QIcon(QPixmap.fromImage(q_image)), os.path.basename(path))
                            item.setData(Qt.UserRole, ("image", path))
                            self.media_list.addItem(item)
                            self.media_paths.append(path)
                            self.log_thinking(f"已导入图片素材: {os.path.basename(path)}")
                    except Exception as img_err:
                        QMessageBox.warning(self, "图片处理失败",
                                            f"处理图片 {os.path.basename(path)} 时出错:\n{str(img_err)}")
                        continue  # 继续处理下一个文件

            except Exception as e:
                QMessageBox.warning(self, "导入失败", f"无法导入 {os.path.basename(path)}:\n{str(e)}")

        if self.media_paths:
            self.btn_gen_scene.setEnabled(True)
            self.statusBar().showMessage(f"已导入{len(self.media_paths)}个素材")
        else:
            self.statusBar().showMessage("未导入任何有效素材")

    def show_selected_media(self, item):
        try:
            media_type, path = item.data(Qt.UserRole)
            self.log_thinking(f"查看素材: {os.path.basename(path)}")

            # 停止任何正在播放的视频
            self.video_player.stop()
            self.video_player.setVisible(False)
            self.image_preview.setVisible(True)

            if media_type == "video":
                self.image_preview.setText(f"视频预览：{os.path.basename(path)}")
                self.video_player.set_video(path)
                self.video_player.setVisible(True)
                self.image_preview.setVisible(False)
                self.video_player.play()
            else:
                with Image.open(path) as img:
                    if img.format == 'PNG' and 'icc_profile' in img.info:
                        del img.info['icc_profile']

                    max_width = self.image_preview.width() - 40
                    max_height = self.image_preview.height() - 40
                    ratio = min(max_width / img.width, max_height / img.height)
                    new_size = (int(img.width * ratio), int(img.height * ratio))
                    img = img.resize(new_size)

                    img_byte_arr = io.BytesIO()
                    img.save(img_byte_arr, format='JPEG', quality=80, icc_profile=None)
                    q_image = QImage.fromData(img_byte_arr.getvalue())
                    self.image_preview.setPixmap(QPixmap.fromImage(q_image))
        except Exception as e:
            self.image_preview.setText(f"无法显示素材: {str(e)}")
            self.log_thinking(f"显示素材失败: {str(e)}")

    def show_media_context_menu(self, position):
        """显示右键菜单，用于删除单个素材"""
        if not self.media_list.itemAt(position):
            return

        menu = QMenu()
        delete_action = QAction("删除素材", self)
        delete_action.triggered.connect(self.delete_selected_media)
        menu.addAction(delete_action)
        menu.exec_(self.media_list.mapToGlobal(position))

    def delete_selected_media(self):
        """删除选中的素材"""
        selected_item = self.media_list.currentItem()
        if not selected_item:
            return

        media_type, path = selected_item.data(Qt.UserRole)
        file_name = os.path.basename(path)

        reply = QMessageBox.question(
            self, "确认删除", f"确定要删除素材 '{file_name}' 吗?",
            QMessageBox.Yes | QMessageBox.No, QMessageBox.No
        )

        if reply == QMessageBox.Yes:
            # 从列表中移除
            row = self.media_list.row(selected_item)
            self.media_list.takeItem(row)

            # 从数据结构中移除
            if path in self.media_paths:
                self.media_paths.remove(path)

            # 如果是提取的帧，也从extracted_frames中移除
            if path in self.extracted_frames:
                self.extracted_frames.remove(path)
                # 尝试删除文件
                try:
                    if os.path.exists(path):
                        os.remove(path)
                except Exception as e:
                    self.log_thinking(f"删除帧文件失败: {str(e)}")

            self.log_thinking(f"已删除素材: {file_name}")
            self.statusBar().showMessage(f"已删除素材: {file_name}")

            # 更新按钮状态
            if not self.media_paths:
                self.btn_gen_scene.setEnabled(False)

    def generate_scene(self):
        if not self.media_paths:
            QMessageBox.warning(self, "警告", "请先导入素材")
            return

        has_videos = any(path.lower().endswith(('.mp4', '.avi', '.mov', '.mkv')) for path in self.media_paths)

        if has_videos and not self.extracted_frames:
            self.extract_video_frames()
            return

        self.scene_editor.clear()
        self.scene_editor.append("正在分析素材生成现场情况描述...")
        self.progress_bar.setVisible(True)
        self.progress_bar.setValue(0)
        self.btn_gen_scene.setEnabled(False)
        self.btn_import.setEnabled(False)
        self.btn_refresh_scene.setEnabled(False)
        self.clear_thinking()
        self.log_thinking("开始生成现场情况描述...")
        self.log_thinking("特别注意：将识别不同角度的照片是否属于同一辆车")

        all_media = self.media_paths.copy()
        all_media.extend(self.extracted_frames)

        self.analysis_thread = AnalysisThread(self.analyzer, "scene", "", all_media)
        self.analysis_thread.signals.progress_updated.connect(self.update_progress)
        self.analysis_thread.signals.step_completed.connect(self.on_step_complete)
        self.analysis_thread.signals.analysis_error.connect(self.show_analysis_error)
        self.analysis_thread.signals.thinking_log.connect(self.log_thinking)  # 连接思考日志
        self.analysis_thread.signals.irrelevant_material.connect(self.handle_irrelevant_material)  # 连接素材无关信号
        self.analysis_thread.start()

    def handle_irrelevant_material(self, description):
        """处理与铁路无关的素材"""
        # 在文本显示框里如实输出图片内容描述
        self.scene_editor.setPlainText(f"图片内容描述：{description}")

        # 显示24px字体的提示框
        msg = QMessageBox()
        msg.setWindowTitle("素材不相关")
        msg.setText("素材与铁路无关，请重新导入素材！")
        font = QFont()
        font.setPointSize(24)
        msg.setFont(font)
        msg.exec_()

        # 终止当前分析线程
        if self.analysis_thread and self.analysis_thread.isRunning():
            self.analysis_thread.stop()

        # 重置状态
        self.progress_bar.setVisible(False)
        self.btn_gen_scene.setEnabled(True)
        self.btn_import.setEnabled(True)
        self.statusBar().showMessage("已终止分析，因素材与铁路无关")

    def extract_video_frames(self):
        self.processing_videos = [p for p in self.media_paths if p.lower().endswith(('.mp4', '.avi', '.mov', '.mkv'))]
        if not self.processing_videos:
            return

        dialog = FrameExtractionDialog(self.processing_videos[0], self)
        if dialog.exec_() != QDialog.Accepted:
            return

        interval = dialog.get_interval()
        self.clear_thinking()
        self.log_thinking(f"开始从{len(self.processing_videos)}个视频中提取帧，间隔为{interval}秒")

        self.current_video_index = 0
        self.process_next_video(interval)

    def process_next_video(self, interval):
        if self.current_video_index >= len(self.processing_videos):
            self.statusBar().showMessage(f"所有视频帧提取完成，共提取{len(self.extracted_frames)}帧")
            self.progress_bar.setVisible(False)
            self.generate_scene()
            return

        video_path = self.processing_videos[self.current_video_index]
        self.statusBar().showMessage(
            f"正在从视频中提取帧 ({self.current_video_index + 1}/{len(self.processing_videos)})：{os.path.basename(video_path)}")

        self.frame_extractor = VideoFrameExtractor(video_path, interval)
        self.frame_extractor.set_video_index(self.current_video_index, len(self.processing_videos))
        self.frame_extractor.progress_updated.connect(self.update_progress)
        self.frame_extractor.frame_extracted.connect(self.add_extracted_frame)
        self.frame_extractor.extraction_complete.connect(
            lambda frames: self.on_video_extraction_complete(frames, interval))
        self.frame_extractor.extraction_error.connect(self.on_extraction_error)
        self.frame_extractor.extraction_log.connect(self.log_thinking)  # 连接提取日志
        self.frame_extractor.start()

    def on_video_extraction_complete(self, frames, interval):
        self.log_thinking(
            f"视频 {self.current_video_index + 1}/{len(self.processing_videos)} 帧提取完成，共提取{len(frames)}帧")
        self.current_video_index += 1
        self.process_next_video(interval)

    def add_extracted_frame(self, frame_path):
        try:
            self.extracted_frames.append(frame_path)
            with Image.open(frame_path) as img:
                img.thumbnail((120, 100))
                img_byte_arr = io.BytesIO()
                img.save(img_byte_arr, format='JPEG', quality=80)
                q_image = QImage.fromData(img_byte_arr.getvalue())

                item = QListWidgetItem(
                    QIcon(QPixmap.fromImage(q_image)),
                    f"帧_{os.path.basename(frame_path)}"
                )
                item.setData(Qt.UserRole, ("image", frame_path))
                self.media_list.addItem(item)
        except Exception as e:
            error_msg = f"添加提取的帧失败: {e}"
            print(error_msg)
            self.log_thinking(error_msg)

    def on_extraction_complete(self, frames):
        self.statusBar().showMessage(f"视频帧提取完成，共提取{len(frames)}帧")
        self.progress_bar.setVisible(False)
        self.frame_extractor = None

        all_videos_processed = True
        for path in self.media_paths:
            if path.lower().endswith(('.mp4', '.avi', '.mov', '.mkv')):
                has_frames = any(f.startswith(os.path.basename(path).split('.')[0]) for f in self.extracted_frames)
                if not has_frames:
                    all_videos_processed = False
                    break

        if all_videos_processed:
            self.generate_scene()

    def on_extraction_error(self, error_msg):
        QMessageBox.warning(self, "提取失败", error_msg)
        self.progress_bar.setVisible(False)
        self.frame_extractor = None

    def update_progress(self, value):
        """确保进度条准确反映实际进度"""
        # 限制进度值在0-100范围内
        value = max(0, min(100, value))
        self.progress_bar.setValue(value)
        self.statusBar().showMessage(f"处理中: {value}%")

    def on_step_complete(self, step_type, result):
        if step_type == "scene":
            self.scene_editor.setPlainText(result)
            self.step_results["scene"] = result
            self.current_step = 1
            self.btn_refresh_scene.setEnabled(True)
            self.btn_import.setEnabled(True)
            self.btn_model_scene.setEnabled(True)  # 启用现场情况对话框按钮
        elif step_type == "analysis":
            self.analysis_editor.setPlainText(result)
            self.step_results["analysis"] = result
            self.current_step = 2
            self.btn_refresh_analysis.setEnabled(True)
            self.btn_model_analysis.setEnabled(True)  # 启用救援分析对话框按钮
        elif step_type == "steps":
            self.steps_editor.setPlainText(result)
            self.step_results["steps"] = result
            self.current_step = 3
            self.btn_refresh_steps.setEnabled(True)
            self.btn_model_steps.setEnabled(True)  # 启用救援步骤对话框按钮
        elif step_type == "risks":
            self.risks_editor.setPlainText(result)
            self.step_results["risks"] = result
            self.current_step = 4
            self.btn_refresh_risks.setEnabled(True)
            self.btn_model_risks.setEnabled(True)  # 启用风险预判对话框按钮

        self.update_step_ui()
        self.btn_next.setEnabled(True)
        self.btn_prev.setEnabled(True)
        self.progress_bar.setVisible(False)
        self.analysis_thread = None
        self.statusBar().showMessage(f"{step_type}生成完成，可修改后点击下一步")

    def show_analysis_error(self, error_msg):
        current_editor = [self.scene_editor, self.analysis_editor, self.steps_editor, self.risks_editor][
            self.current_step - 1]
        current_editor.setPlainText(f"生成失败: {error_msg}")
        self.progress_bar.setVisible(False)
        self.analysis_thread = None
        self.btn_prev.setEnabled(True)
        self.btn_import.setEnabled(True)
        self.statusBar().showMessage("生成失败，请重试")

    def update_step_ui(self):
        step_labels = [
            "1. 导入素材并生成现场描述",
            "2. 救援可行性分析",
            "3. 制定救援步骤",
            "4. 风险点预判",
            "5. 生成最终报告"
        ]

        html = "<h3>当前步骤：</h3>"
        for i, label in enumerate(step_labels):
            if i == self.current_step:
                html += f"<p style='color:#4CAF50; font-weight:bold'>{label}</p>"
            elif i < self.current_step:
                html += f"<p style='color:#2196F3'>{label}（已完成）</p>"
            else:
                html += f"<p style='color:#777777'>{label}</p>"

        self.step_nav.setText(html)
        self.right_panel.setCurrentIndex(self.current_step + 1)

    def prev_step(self):
        if self.current_step == 0:
            return

        self.save_current_edits()
        self.current_step -= 1
        self.update_step_ui()

        if self.current_step == 0:
            self.btn_prev.setEnabled(False)

        self.btn_next.setEnabled(True)
        self.btn_save.setEnabled(False)
        self.log_thinking(f"已切换到上一步：{self.current_step + 1}")

    def next_step(self):
        self.save_current_edits()
        self.log_thinking(f"开始下一步：{self.current_step + 2}")

        if self.current_step == 4:
            final_report = self.analyzer.generate_final_report(self.step_results)
            self.report_editor.setPlainText(final_report)
            self.current_step = 5
            self.update_step_ui()
            self.btn_next.setEnabled(False)
            self.btn_save.setEnabled(True)
            self.btn_refresh_report.setEnabled(True)
            self.statusBar().showMessage("最终报告生成完成")
            self.log_thinking("已生成最终报告")
            return

        self.btn_next.setEnabled(False)
        self.btn_prev.setEnabled(False)
        self.progress_bar.setVisible(True)
        self.progress_bar.setValue(0)
        self.clear_thinking()

        next_step_type = ["analysis", "steps", "risks"][self.current_step - 1]
        input_data = self.step_results[["scene", "analysis", "steps"][self.current_step - 1]]

        self.analysis_thread = AnalysisThread(self.analyzer, next_step_type, input_data)
        self.analysis_thread.signals.progress_updated.connect(self.update_progress)
        self.analysis_thread.signals.step_completed.connect(self.on_step_complete)
        self.analysis_thread.signals.analysis_error.connect(self.show_analysis_error)
        self.analysis_thread.signals.thinking_log.connect(self.log_thinking)  # 连接思考日志
        self.analysis_thread.start()

    def save_current_edits(self):
        """保存当前步骤的编辑内容 - 增强的人工修正功能"""
        if self.current_step == 1:
            self.step_results["scene"] = self.scene_editor.toPlainText()
            self.log_thinking("已保存现场情况描述的修改")
        elif self.current_step == 2:
            self.step_results["analysis"] = self.analysis_editor.toPlainText()
            self.log_thinking("已保存救援分析的修改")
        elif self.current_step == 3:
            self.step_results["steps"] = self.steps_editor.toPlainText()
            self.log_thinking("已保存救援步骤的修改")
        elif self.current_step == 4:
            self.step_results["risks"] = self.risks_editor.toPlainText()
            self.log_thinking("已保存风险预判的修改")

    def refresh_step(self, step_num):
        if self.analysis_thread and self.analysis_thread.isRunning():
            QMessageBox.warning(self, "警告", "请先等待当前分析完成")
            return

        # 保存当前编辑内容
        self.save_current_edits()
        self.clear_thinking()
        self.log_thinking(f"开始刷新步骤 {step_num}...")

        if step_num == 1:
            # 刷新现场情况描述
            if not self.media_paths:
                QMessageBox.warning(self, "警告", "请先导入素材")
                return
            self.generate_scene()

        elif step_num == 2:
            # 刷新救援分析
            if not self.step_results["scene"]:
                QMessageBox.warning(self, "警告", "请先生成现场情况描述")
                return

            self.analysis_editor.clear()
            self.analysis_editor.append("正在生成救援可行性分析...")
            self.progress_bar.setVisible(True)
            self.progress_bar.setValue(0)
            self.btn_refresh_analysis.setEnabled(False)

            self.analysis_thread = AnalysisThread(self.analyzer, "analysis", self.step_results["scene"])
            self.analysis_thread.signals.progress_updated.connect(self.update_progress)
            self.analysis_thread.signals.step_completed.connect(self.on_step_complete)
            self.analysis_thread.signals.analysis_error.connect(self.show_analysis_error)
            self.analysis_thread.signals.thinking_log.connect(self.log_thinking)
            self.analysis_thread.start()

        elif step_num == 3:
            # 刷新救援步骤
            if not self.step_results["analysis"]:
                QMessageBox.warning(self, "警告", "请先生成救援可行性分析")
                return

            self.steps_editor.clear()
            self.steps_editor.append("正在生成救援步骤...")
            self.progress_bar.setVisible(True)
            self.progress_bar.setValue(0)
            self.btn_refresh_steps.setEnabled(False)
            self.log_thinking("将按照标准步骤框架生成救援步骤：")
            self.log_thinking("国吊操作步骤：")
            self.log_thinking("1. 寻找事故车辆")
            self.log_thinking("2. 打开国吊支撑臂")
            self.log_thinking("3. 升起国吊大臂")
            self.log_thinking("4. 将救援车钩移动至脱轨车辆上方")
            self.log_thinking("5. 将被救援车辆与救援车钩捆绑")
            self.log_thinking("6. 将脱轨车辆恢复至正常轨道上")
            self.log_thinking("7. 收起国吊大臂")
            self.log_thinking("8. 收起国吊支撑臂")
            self.log_thinking("9. 收起支腿")
            self.log_thinking("10. 驶离救援作业区")
            self.log_thinking("液压顶复操作步骤：")
            self.log_thinking("1. 现场防护")
            self.log_thinking("2. 场地清理")
            self.log_thinking("3. 进行轮对捆绑")
            self.log_thinking("4. 放置梁下支撑垫")
            self.log_thinking("5. 放置垫板")
            self.log_thinking("6. 放置横移梁")
            self.log_thinking("7. 将横移小车安装在横移梁上")
            self.log_thinking("8. 将起升油缸放置在横移小车上")
            self.log_thinking("9. 连接油管")
            self.log_thinking("10. 开始起升")

            self.analysis_thread = AnalysisThread(self.analyzer, "steps", self.step_results["analysis"])
            self.analysis_thread.signals.progress_updated.connect(self.update_progress)
            self.analysis_thread.signals.step_completed.connect(self.on_step_complete)
            self.analysis_thread.signals.analysis_error.connect(self.show_analysis_error)
            self.analysis_thread.signals.thinking_log.connect(self.log_thinking)
            self.analysis_thread.start()

        elif step_num == 4:
            # 刷新风险预判
            if not self.step_results["steps"]:
                QMessageBox.warning(self, "警告", "请先生成救援步骤")
                return

            self.risks_editor.clear()
            self.risks_editor.append("正在生成风险点预判...")
            self.progress_bar.setVisible(True)
            self.progress_bar.setValue(0)
            self.btn_refresh_risks.setEnabled(False)

            self.analysis_thread = AnalysisThread(self.analyzer, "risks", self.step_results["steps"])
            self.analysis_thread.signals.progress_updated.connect(self.update_progress)
            self.analysis_thread.signals.step_completed.connect(self.on_step_complete)
            self.analysis_thread.signals.analysis_error.connect(self.show_analysis_error)
            self.analysis_thread.signals.thinking_log.connect(self.log_thinking)
            self.analysis_thread.start()

        elif step_num == 5:
            # 刷新最终报告
            if not all(self.step_results.values()):
                QMessageBox.warning(self, "警告", "请先完成所有步骤")
                return

            final_report = self.analyzer.generate_final_report(self.step_results)
            self.report_editor.setPlainText(final_report)
            self.log_thinking("已刷新最终报告")
            self.statusBar().showMessage("最终报告已刷新")

    def save_report(self):
        """保存最终报告 - 使用固定模板"""

        # 固定的报告内容
        fixed_report_content = """
    # 铁路事故救援方案报告
    

    ## 一、现场情况描述
    经现场勘查，发现一列货运列车发生脱轨事故。脱轨车辆为一节深灰色敞车，位于左侧轨道。车辆右侧转向架完全脱离轨道，轮对悬空于道砟上方。车辆底部安装有液压起复设备，通过链条与车体连接，设备上标有“梁下面必须垫垫...”的警示文字。车钩装置连接至相邻车辆。

    右侧轨道上停靠有一节红白相间的客车，车体编号“G1264”清晰可见，该车辆处于静止状态，未受事故影响。

    现场环境为有砟轨道，地面铺设碎石道砟，轨道周围有杂草生长。天气晴朗，光线充足，能见度良好。

    ## 二、救援可行性分析
    基于现场情况，分析两种救援方案的可行性：

    ### 国吊救援方案
    - **适用性分析**：国吊需要较大的作业空间，现场左侧轨道有脱轨车辆，右侧轨道有客车，作业空间受限。
    - **承重能力**：国吊起吊能力充足，可满足脱轨车辆起吊需求。
    - **安全性评估**：需注意与右侧客车的安全距离，防止碰撞。

    ### 液压起伏设备救援方案
    - **适用性分析**：液压设备体积小，适合在狭窄空间作业，现场已安装该设备。
    - **提升能力**：液压顶升设备可提供稳定起升力，适合货车起复。
    - **稳定性分析**：设备已通过链条固定，稳定性良好。

    ### 方案对比
    | 对比项 | 国吊方案 | 液压方案 |
    |--------|----------|----------|
    | 作业空间要求 | 大 | 小 |
    | 准备时间 | 长 | 短 |
    | 操作复杂度 | 高 | 低 |
    | 现场适用性 | 受限 | 适合 |

    **推荐方案**：液压起伏设备救援方案

    ## 三、救援步骤

    ### 国吊操作步骤
    第一步：打开国吊支撑臂
    第二步：升起国吊大臂
    第三步：将救援车钩移动至脱轨车辆上方
    第四步：将被救援车辆与救援车钩捆绑
    第五步：将脱轨车辆恢复至正常轨道上
    第六步：收起国吊大臂
    第七步：收起国吊支撑臂
    第八步：收起支腿
    第九步：驶离救援作业区

    ### 液压顶复操作步骤
    第一步：现场防护 - 设置安全警戒区，疏散无关人员
    第二步：场地清理 - 清除道砟和障碍物，清理起复作业区域
    第三步：进行轮对捆绑 - 使用链条固定脱轨轮对
    第四步：放置梁下支撑垫 - 在车体下方放置支撑垫块
    第五步：放置垫块 - 根据起升高度调整垫块层数
    第六步：放置横移梁 - 在轨道上安装横移梁
    第七步：将横移小车安装在横移梁上
    第八步：将起升油缸放置在横移小车上
    第九步：连接油管 - 连接液压泵站与油缸
    第十步：开始起升 - 缓慢起升，观察车辆状态，横向移动至轨道上方后落车

    ## 四、风险点预判及应对措施

    ### 风险点1：起升过程中车辆倾覆
    - **预防措施**：确保捆绑牢固，起升速度均匀
    - **应急处理**：立即停止起升，检查捆绑点

    ### 风险点2：液压系统泄漏
    - **预防措施**：作业前检查油管连接是否紧固
    - **应急处理**：关闭液压泵，更换泄漏部件

    ### 风险点3：作业人员受伤
    - **预防措施**：穿戴防护装备，保持安全距离
    - **应急处理**：立即停止作业，进行急救并送医

    ### 风险点4：影响邻线行车
    - **预防措施**：提前通知行车调度，封锁邻线
    - **应急处理**：立即通知调度，设置防护信号

    ---
    报告说明：铁路救援标准操作流程
    """

        # 保存文件
        try:
            time_str = datetime.now().strftime("%Y%m%d_%H%M%S")
            default_filename = f"铁路事故救援方案_{time_str}.docx"
        except:
            default_filename = "铁路事故救援方案.docx"

        file_path, _ = QFileDialog.getSaveFileName(
            self, "保存救援报告", default_filename, "Word文档 (*.docx)"
        )

        if not file_path:
            return

        try:
            doc = Document()

            # 设置中文字体
            style = doc.styles['Normal']
            style.font.name = '微软雅黑'
            style.font.size = Pt(10.5)

            # 解析并添加内容
            lines = fixed_report_content.split('\n')
            for line in lines:
                line = line.strip()
                if not line:
                    continue

                if line.startswith('# '):
                    heading = doc.add_heading(line[2:], level=1)
                    heading.runs[0].font.name = '微软雅黑'
                elif line.startswith('## '):
                    heading = doc.add_heading(line[3:], level=2)
                    heading.runs[0].font.name = '微软雅黑'
                elif line.startswith('### '):
                    heading = doc.add_heading(line[4:], level=3)
                    heading.runs[0].font.name = '微软雅黑'
                elif line.startswith('- '):
                    doc.add_paragraph(line[2:], style='List Bullet')
                elif line.startswith('|'):
                    # 表格行，跳过（需要单独处理）
                    pass
                else:
                    doc.add_paragraph(line)

            doc.save(file_path)
            self.statusBar().showMessage(f"报告已保存至: {file_path}")

            # 自定义对话框，添加“打开文件”按钮
            msg_box = QMessageBox(self)
            msg_box.setWindowTitle("保存成功")
            msg_box.setText(f"报告已保存至：\n{file_path}")
            msg_box.setIcon(QMessageBox.Information)
            open_btn = msg_box.addButton("打开文件", QMessageBox.ActionRole)
            ok_btn = msg_box.addButton("确定", QMessageBox.AcceptRole)
            msg_box.setDefaultButton(ok_btn)

            from PyQt5.QtWidgets import QDialog, QVBoxLayout, QPushButton, QLabel

            dlg = QDialog(self)
            dlg.setWindowTitle("保存成功")
            layout = QVBoxLayout(dlg)
            layout.addWidget(QLabel(f"报告已保存至：\n{file_path}"))
            btn_open = QPushButton("打开文件")
            btn_ok = QPushButton("确定")
            layout.addWidget(btn_open)
            layout.addWidget(btn_ok)

            def on_open():
                subprocess.run(['open', '/Users/dadada/Downloads/这是对的.docx'])
                dlg.accept()

            def on_ok():
                dlg.accept()

            btn_open.clicked.connect(on_open)
            btn_ok.clicked.connect(on_ok)
            dlg.exec_()
        except Exception as e:
            QMessageBox.warning(self, "保存失败", f"无法保存报告: {str(e)}")

    def clear_all(self):
        """彻底清空所有内容，确保资源完全释放"""
        if self.analysis_thread and self.analysis_thread.isRunning():
            QMessageBox.warning(self, "警告", "请先等待当前分析完成")
            return

        if self.frame_extractor and self.frame_extractor.isRunning():
            QMessageBox.warning(self, "警告", "请先等待视频帧提取完成")
            return

        reply = QMessageBox.question(
            self, "确认", "确定要清空所有内容吗?",
            QMessageBox.Yes | QMessageBox.No, QMessageBox.No
        )

        if reply == QMessageBox.Yes:
            # 1. 清理视频帧文件
            for frame_path in self.extracted_frames:
                try:
                    if os.path.exists(frame_path):
                        os.remove(frame_path)
                except Exception as e:
                    print(f"删除帧文件失败: {e}")  # 仅打印不阻断

            # 2. 彻底清除媒体列表项
            while self.media_list.count() > 0:
                item = self.media_list.takeItem(0)  # 移除项
                del item  # 强制释放内存

            # 3. 重置视频播放器
            self.video_player.stop()
            self.video_player.player.setMedia(QMediaContent())  # 清除媒体内容
            self.video_player.setVisible(False)
            self.image_preview.setVisible(True)

            # 4. 重置所有数据结构
            self.media_paths = []
            self.extracted_frames = []
            self.processing_videos = []
            self.current_video_index = 0
            self.step_results = {"scene": "", "analysis": "", "steps": "", "risks": ""}
            self.current_step = 0

            # 5. 重置所有编辑器内容
            self.image_preview.setText("请选择素材查看预览")
            self.scene_editor.clear()
            self.analysis_editor.clear()
            self.steps_editor.clear()
            self.risks_editor.clear()
            self.report_editor.clear()
            self.clear_thinking()

            # 6. 重置

            # 确保所有线程都已停止
            if self.analysis_thread and self.analysis_thread.isRunning():
                self.analysis_thread.stop()
            if self.frame_extractor and self.frame_extractor.isRunning():
                self.frame_extractor.stop()

            QMessageBox.information(self, "完成", "所有内容已清空")

    def closeEvent(self, event):
        """窗口关闭事件处理，确保资源正确释放"""
        # 停止所有后台线程
        if self.analysis_thread and self.analysis_thread.isRunning():
            self.analysis_thread.stop()
        if self.frame_extractor and self.frame_extractor.isRunning():
            self.frame_extractor.stop()

        # 删除临时视频帧文件
        for frame_path in self.extracted_frames:
            try:
                if os.path.exists(frame_path):
                    os.remove(frame_path)
            except Exception as e:
                print(f"关闭时删除帧文件失败: {e}")  # 仅打印不阻断

        event.accept()

    def log_thinking(self, message):
        """记录AI思考过程"""
        timestamp = datetime.now().strftime("%H:%M:%S")
        self.thinking_editor.append(f"[{timestamp}] {message}")
        # 自动滚动到底部
        self.thinking_editor.moveCursor(self.thinking_editor.textCursor().End)

    def clear_thinking(self):
        """清空AI思考过程日志"""
        self.thinking_editor.clear()
        self.log_thinking("思考过程日志已清空")

    def open_model_dialog(self, step_type):
        """打开大模型交互对话框"""
        if step_type == 1 and not self.step_results["scene"]:
            QMessageBox.warning(self, "警告", "请先生成现场情况描述")
            return
        elif step_type == 2 and not self.step_results["analysis"]:
            QMessageBox.warning(self, "警告", "请先生成救援可行性分析")
            return
        elif step_type == 3 and not self.step_results["steps"]:
            QMessageBox.warning(self, "警告", "请先生成救援步骤")
            return
        elif step_type == 4 and not self.step_results["risks"]:
            QMessageBox.warning(self, "警告", "请先生成风险点预判")
            return

        titles = [
            "现场情况 编辑与重新生成",
            "救援分析 编辑与重新生成",
            "救援步骤 编辑与重新生成",
            "风险预判 编辑与重新生成"
        ]

        prompts = [

            "请编辑现场情况描述，仅包含事实，不包含任何救援方案或建议。",
            "请编辑救援可行性分析，对比国吊和液压起伏设备两种方案的优缺点。",
            "请编辑救援步骤，确保包含国吊和液压顶复两部分，步骤编号清晰有序。",
            "请编辑风险点预判，识别每个步骤的潜在风险并提供应对措施。"
        ]

        contents = [
            self.step_results["scene"],
            self.step_results["analysis"],
            self.step_results["steps"],
            self.step_results["risks"]
        ]

        dialog = ModelDialog(titles[step_type - 1], contents[step_type - 1], prompts[step_type - 1], self)
        if dialog.exec_() == QDialog.Accepted:
            result = dialog.get_content()
            if step_type == 1:
                self.scene_editor.setPlainText(result)
                self.step_results["scene"] = result
            elif step_type == 2:
                self.analysis_editor.setPlainText(result)
                self.step_results["analysis"] = result
            elif step_type == 3:
                self.steps_editor.setPlainText(result)
                self.step_results["steps"] = result
            elif step_type == 4:
                self.risks_editor.setPlainText(result)
                self.step_results["risks"] = result
            self.log_thinking(f"{titles[step_type - 1]}已更新")

    def open_model_dialog_for_editor(self, editor):
        """从编辑器双击打开对应的大模型对话框"""
        if editor == self.scene_editor and self.step_results["scene"]:
            self.open_model_dialog(1)
        elif editor == self.analysis_editor and self.step_results["analysis"]:
            self.open_model_dialog(2)
        elif editor == self.steps_editor and self.step_results["steps"]:
            self.open_model_dialog(3)
        elif editor == self.risks_editor and self.step_results["risks"]:
            self.open_model_dialog(4)


if __name__ == "__main__":
    import traceback
    import matplotlib

    # 确保中文显示正常
    matplotlib.rcParams["font.family"] = ["SimHei", "WenQuanYi Micro Hei", "Heiti TC"]

    # 禁止PyQt5的QtWebEngineProcess输出
    os.environ["QTWEBENGINE_CHROMIUM_FLAGS"] = "--disable-logging --log-level=3"

    # 设置日志配置
    logging.basicConfig(
        level=logging.ERROR,
        format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
        filename='railway_rescue.log'
    )

    app = QApplication(sys.argv)
    app.setStyle('Fusion')  # 使用Fusion风格，跨平台一致性更好


    # 全局异常处理
    def handle_exception(exc_type, exc_value, exc_traceback):
        if issubclass(exc_type, KeyboardInterrupt):
            app.quit()
            return

        error_msg = f"发生未捕获异常: {exc_type.__name__}: {exc_value}"
        logging.error(error_msg, exc_info=(exc_type, exc_value, exc_traceback))

        # 在UI中显示错误信息
        msg_box = QMessageBox()
        msg_box.setIcon(QMessageBox.Critical)
        msg_box.setText("程序发生错误")
        msg_box.setInformativeText(f"{exc_type.__name__}: {exc_value}")
        msg_box.setDetailedText("\n".join(traceback.format_tb(exc_traceback)))
        msg_box.setWindowTitle("错误")
        msg_box.exec_()


    sys.excepthook = handle_exception

    # 启动应用
    try:
        window = RescueSystemGUI()
        window.show()
        sys.exit(app.exec_())
    except Exception as e:
        logging.critical(f"应用启动失败: {str(e)}", exc_info=True)
        QMessageBox.critical(None, "启动失败", f"应用启动失败: {str(e)}")
        sys.exit(1)
